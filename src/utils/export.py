import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
from bs4 import BeautifulSoup

def markdown_to_docx(markdown_text: str) -> bytes:
    """Convert markdown text to a formatted Word document."""
    
    # Convert markdown to HTML
    html = markdown.markdown(markdown_text)
    soup = BeautifulSoup(html, 'html.parser')
    
    # Create document
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    def process_element(element, level=0):
        """Process HTML elements and apply appropriate Word formatting."""
        if element.name == 'h1':
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(element.get_text())
            run.bold = True
            run.font.size = Pt(18)
            paragraph.space_after = Pt(12)
        
        elif element.name == 'h2':
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(element.get_text())
            run.bold = True
            run.font.size = Pt(16)
            paragraph.space_after = Pt(10)
        
        elif element.name == 'h3':
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(element.get_text())
            run.bold = True
            run.font.size = Pt(14)
            paragraph.space_after = Pt(8)
        
        elif element.name == 'p':
            paragraph = doc.add_paragraph()
            paragraph.add_run(element.get_text())
            paragraph.space_after = Pt(8)
        
        elif element.name == 'ul':
            for li in element.find_all('li', recursive=False):
                paragraph = doc.add_paragraph()
                paragraph.style = 'List Bullet'
                paragraph.add_run(li.get_text())
                paragraph.space_after = Pt(4)
        
        elif element.name == 'ol':
            for i, li in enumerate(element.find_all('li', recursive=False), 1):
                paragraph = doc.add_paragraph()
                paragraph.style = 'List Number'
                paragraph.add_run(li.get_text())
                paragraph.space_after = Pt(4)
    
    # Process all elements
    for element in soup.children:
        if element.name:
            process_element(element)
    
    # Save to bytes
    docx_file = io.BytesIO()
    doc.save(docx_file)
    docx_file.seek(0)
    
    return docx_file.getvalue() 